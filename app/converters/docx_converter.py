from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from bs4 import BeautifulSoup
from app.utils.file_utils import make_output_path
import csv


def docx_to_txt(input_path: str) -> str:
    """Extract all text from DOCX and save as .txt"""
    output_path = make_output_path(input_path, "txt")
    doc = Document(input_path)
    lines = [para.text for para in doc.paragraphs]
    with open(output_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    return output_path


def docx_to_pdf(input_path: str) -> str:
    """Convert DOCX text content to PDF."""
    output_path = make_output_path(input_path, "pdf")
    doc = Document(input_path)
    styles = getSampleStyleSheet()
    story = []
    for para in doc.paragraphs:
        if para.text.strip():
            story.append(Paragraph(para.text, styles["Normal"]))
            story.append(Spacer(1, 6))
    pdf = SimpleDocTemplate(output_path, pagesize=letter)
    pdf.build(story)
    return output_path


def docx_to_html(input_path: str) -> str:
    """Convert DOCX paragraphs to a clean HTML file."""
    output_path = make_output_path(input_path, "html")
    doc = Document(input_path)
    paragraphs_html = "\n".join(
        f"<p>{para.text}</p>"
        for para in doc.paragraphs if para.text.strip()
    )
    html = f"""<!DOCTYPE html>
<html>
<head><meta charset='utf-8'><title>Converted Document</title></head>
<body>
{paragraphs_html}
</body>
</html>"""
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(html)
    return output_path


def docx_to_csv(input_path: str) -> str:
    """Extract tables from DOCX and save as CSV."""
    output_path = make_output_path(input_path, "csv")
    doc = Document(input_path)
    with open(output_path, "w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        for table in doc.tables:
            for row in table.rows:
                writer.writerow([cell.text.strip() for cell in row.cells])
    return output_path