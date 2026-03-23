from bs4 import BeautifulSoup
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
from app.utils.file_utils import make_output_path


def html_to_txt(input_path: str) -> str:
    """Strip HTML tags and save plain text."""
    output_path = make_output_path(input_path, "txt")
    with open(input_path, encoding="utf-8") as f:
        soup = BeautifulSoup(f.read(), "lxml")
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(soup.get_text(separator="\n"))
    return output_path


def html_to_pdf(input_path: str) -> str:
    """Convert HTML content to PDF."""
    output_path = make_output_path(input_path, "pdf")
    with open(input_path, encoding="utf-8") as f:
        soup = BeautifulSoup(f.read(), "lxml")
    text_blocks = [p.get_text() for p in soup.find_all(["p", "h1", "h2", "h3", "li"])]
    styles = getSampleStyleSheet()
    story = []
    for block in text_blocks:
        if block.strip():
            story.append(Paragraph(block.strip(), styles["Normal"]))
            story.append(Spacer(1, 6))
    doc = SimpleDocTemplate(output_path, pagesize=letter)
    doc.build(story)
    return output_path


def html_to_docx(input_path: str) -> str:
    """Convert HTML content to DOCX."""
    output_path = make_output_path(input_path, "docx")
    with open(input_path, encoding="utf-8") as f:
        soup = BeautifulSoup(f.read(), "lxml")
    doc = Document()
    for tag in soup.find_all(["h1", "h2", "h3", "p", "li"]):
        text = tag.get_text().strip()
        if not text:
            continue
        if tag.name == "h1":
            doc.add_heading(text, level=1)
        elif tag.name == "h2":
            doc.add_heading(text, level=2)
        elif tag.name == "h3":
            doc.add_heading(text, level=3)
        else:
            doc.add_paragraph(text)
    doc.save(output_path)
    return output_path